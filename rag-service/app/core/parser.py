"""文档解析器 - 主流 RAG 模式：多格式解析 + 分块。

支持格式：
- PDF: PyMuPDF 文本层提取（毫秒级，电子版 PDF）
- DOCX: python-docx 段落 + 表格
- TXT/MD: 直接读取

输出：Markdown 格式纯文本 + 分块列表。
解析结果缓存到 parsed/{doc_id}.md，供预览和重新索引使用。
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# 主流 RAG 默认分块参数
DEFAULT_CHUNK_SIZE = 512  # 字符数
DEFAULT_CHUNK_OVERLAP = 64


def parse_file(file_path: str | Path, doc_id: str) -> Dict:
    """解析文档，返回 {text, pages, chunks, meta}。

    text: 完整 Markdown 文本（用于预览）
    pages: 页数（PDF 专有，其他为 1）
    chunks: 分块列表 [{id, text, page}]
    meta: {file_type, size_bytes, parser}
    """
    p = Path(file_path)
    ext = p.suffix.lower()
    if not p.exists():
        raise FileNotFoundError(f"file not found: {p}")

    if ext == ".pdf":
        text, pages = _parse_pdf(p)
        parser = "pymupdf"
    elif ext == ".docx":
        text, pages = _parse_docx(p)
        parser = "python-docx"
    elif ext in {".txt", ".md"}:
        text = p.read_text(encoding="utf-8", errors="ignore")
        pages = 1
        parser = "plain"
    elif ext == ".doc":
        raise ValueError(".doc 老格式不支持，请转换为 .docx")
    else:
        # 兜底：尝试按文本读
        try:
            text = p.read_text(encoding="utf-8", errors="ignore")
            pages = 1
            parser = "fallback-plain"
        except Exception as e:
            raise ValueError(f"不支持的文件格式 {ext}: {e}")

    chunks = chunk_text(text, doc_id=doc_id)
    return {
        "text": text,
        "pages": pages,
        "chunks": chunks,
        "meta": {
            "file_type": ext,
            "size_bytes": p.stat().st_size,
            "parser": parser,
            "chunk_count": len(chunks),
        },
    }


def _parse_pdf(p: Path) -> Tuple[str, int]:
    """PyMuPDF 提取 PDF 文本层，输出 Markdown。"""
    import fitz

    doc = fitz.open(str(p))
    pages_text: List[str] = []
    for i, page in enumerate(doc, start=1):
        blocks = page.get_text("blocks")  # [(x0,y0,x1,y1,text,block_no,block_type),...]
        page_lines: List[str] = []
        for b in blocks:
            text = b[4].strip() if len(b) > 4 else ""
            if not text:
                continue
            # 块级分隔
            page_lines.append(text)
        if page_lines:
            pages_text.append(f"## 第 {i} 页\n\n" + "\n\n".join(page_lines))
    doc.close()

    full = "\n\n---\n\n".join(pages_text) if pages_text else ""
    if not full.strip():
        logger.warning("PDF %s 文本层为空（可能是扫描件），建议 OCR", p.name)
    return full, len(pages_text) if pages_text else 0


def _parse_docx(p: Path) -> Tuple[str, int]:
    """python-docx 提取段落和表格，输出 Markdown。"""
    import docx

    d = docx.Document(str(p))
    out: List[str] = []
    for para in d.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = (para.style.name or "").lower()
        if "heading 1" in style or "title" in style:
            out.append(f"# {t}")
        elif "heading 2" in style:
            out.append(f"## {t}")
        elif "heading 3" in style:
            out.append(f"### {t}")
        elif "list" in style:
            out.append(f"- {t}")
        else:
            out.append(t)

    # 表格转 Markdown
    for table in d.tables:
        rows = table.rows
        if not rows:
            continue
        md_rows = []
        for r in rows:
            cells = [c.text.strip().replace("|", "\\|") for c in r.cells]
            md_rows.append("| " + " | ".join(cells) + " |")
        if len(md_rows) >= 1:
            # 表头分隔行
            sep = "| " + " | ".join(["---"] * len(rows[0].cells)) + " |"
            md_rows.insert(1, sep)
            out.append("\n".join(md_rows))

    return "\n\n".join(out), 1


def chunk_text(text: str, doc_id: str = "", chunk_size: int = DEFAULT_CHUNK_SIZE,
               overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[Dict]:
    """按字符数分块（主流 RAG 简单切分），保留重叠。

    优先按段落边界切，段落过长再按字符切。
    """
    if not text or not text.strip():
        return []

    # 按双换行切段落
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: List[Dict] = []
    buf = ""
    buf_start = 0
    idx = 0

    def flush(end: int):
        nonlocal buf, buf_start, idx
        if buf.strip():
            chunks.append({
                "id": f"{doc_id}_{idx}" if doc_id else f"chunk_{idx}",
                "text": buf.strip(),
                "start": buf_start,
                "end": end,
            })
            idx += 1
        buf = ""

    pos = 0
    for para in paras:
        if len(buf) + len(para) + 2 > chunk_size and buf:
            flush(pos)
            # 重叠：保留尾部
            buf = buf[-overlap:] if overlap > 0 else ""
            buf_start = pos - len(buf)
        buf = (buf + "\n\n" + para) if buf else para
        pos += len(para) + 2
    flush(pos)
    return chunks


def save_parsed(parsed_dir: Path, doc_id: str, text: str) -> Path:
    """把解析后文本存到 parsed/{doc_id}.md。"""
    parsed_dir.mkdir(parents=True, exist_ok=True)
    out = parsed_dir / f"{doc_id}.md"
    out.write_text(text, encoding="utf-8")
    return out


def load_parsed(parsed_dir: Path, doc_id: str) -> Optional[str]:
    """读取解析后文本，不存在返回 None。"""
    f = parsed_dir / f"{doc_id}.md"
    return f.read_text(encoding="utf-8", errors="ignore") if f.exists() else None
